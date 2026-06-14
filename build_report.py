# -*- coding: utf-8 -*-
"""
ODEV-2 PDF/Word raporunu, train_and_evaluate.py'nin urettigi GERCEK ciktilardan otomatik olusturur.
Cikti: rapor/Odev2_Raporu.docx  (ve docx2pdf ile Odev2_Raporu.pdf)
Calistirma: py -3.9 build_report.py
"""
import os
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(BASE, "outputs")
REP = os.path.join(BASE, "rapor")
os.makedirs(REP, exist_ok=True)

cos = pd.read_csv(os.path.join(OUT, "cosine_eval.csv"))
sem = pd.read_csv(os.path.join(OUT, "semantic_eval.csv"))
sw = pd.read_csv(os.path.join(OUT, "similar_words.csv"))
top5 = pd.read_csv(os.path.join(OUT, "top5_per_model.csv"))
jac = pd.read_csv(os.path.join(OUT, "jaccard_matrix.csv"), index_col=0)
summary = pd.read_csv(os.path.join(OUT, "summary.csv"))
query_txt = open(os.path.join(OUT, "query.txt"), encoding="utf-8").read().strip().splitlines()

doc = Document()

# Varsayilan yazi tipi
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def H(text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def P(text, bold=False, italic=False, size=11, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def df_to_table(df, header_color="2E5496", col_widths=None, font_size=8):
    t = doc.add_table(rows=1, cols=len(df.columns))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for j, c in enumerate(df.columns):
        hdr[j].text = str(c)
        for pr in hdr[j].paragraphs:
            for rn in pr.runs:
                rn.bold = True
                rn.font.size = Pt(font_size)
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for j, c in enumerate(df.columns):
            cells[j].text = str(row[c])
            for pr in cells[j].paragraphs:
                for rn in pr.runs:
                    rn.font.size = Pt(font_size)
    if col_widths:
        for j, w in enumerate(col_widths):
            for r in t.rows:
                r.cells[j].width = Inches(w)
    return t


# ======================================================================================
# KAPAK
# ======================================================================================
P("YAPAY ZEKA DERSI", bold=True, size=14, align="center")
P("Odev-2", size=12, align="center")
doc.add_paragraph()
P("EGITILEN WORD2VEC MODELLERI ILE", bold=True, size=18, align="center")
P("METIN BENZERLIGI HESAPLAMA VE DEGERLENDIRME", bold=True, size=18, align="center")
doc.add_paragraph()
P("Proje Konusu: Mobil Uygulama (Seyahat Uygulamasi) Yorum Analizi", size=12, align="center")
doc.add_paragraph()
doc.add_paragraph()
P("Proje Sahipleri", bold=True, size=12, align="center")
P("Mirac BIRBEN  -  2202131003", size=12, align="center")
P("Ayhan TUGALAY  -  2202131012", size=12, align="center")
P("Muhammet Ali CEBI  -  2202131019", size=12, align="center")
doc.add_paragraph()
P("YBS 4. Sinif", size=11, align="center")
P("Teslim Tarihi: 15 Haziran 2026", size=11, align="center")
P("GitHub: https://github.com/miracbirbenhub/YAPAYZEKAPROJE", size=10, align="center")
doc.add_page_break()

# ======================================================================================
# 1. GIRIS
# ======================================================================================
H("1. Giris", 1)

H("1.1 Odevin Amaci", 2)
P("Bu odevin amaci, birinci odevde on isleme (pre-processing) tabi tuttugumuz iki temiz veri seti "
  "(lemmatized ve stemmed) uzerinde Gensim kutuphanesi ile Word2Vec modelleri egitmek ve bu "
  "modelleri kullanarak metinler arasi benzerlik hesaplamalari yapmaktir. Toplam 16 model "
  "(2 veri seti x 8 parametre seti) egitilmis; secilen bir ornek giris metnine en cok benzeyen 5 "
  "yorum her model icin ayri ayri bulunmus ve modeller uc farkli yontemle (Cosine, Anlamsal, "
  "Jaccard) karsilastirmali olarak degerlendirilmistir.")

H("1.2 Kullanilan Veri Seti", 2)
P("Proje konumuz bir mobil seyahat uygulamasina (itinerary/trip planlama uygulamasi) yapilan "
  "kullanici yorumlarinin analizidir. Veri seti, farkli platform (Android/iOS) ve ulkelerden gelen "
  "Ingilizce kullanici yorumlarindan olusmaktadir. Her yorum; tarih, platform, ulke, yorum metni, "
  "yildiz (1-5), kullanici kimligi ve etkilesim sayilari gibi alanlar icermektedir. Veri seti "
  "Odev-1 kapsaminda toplanip temizlenmistir.")
bullet("Toplam yorum (satir) sayisi: 321")
bullet("Yildiz dagilimi: 1 yildiz=49, 2=78, 3=70, 4=73, 5=51 (dengeli bir dagilim)")
bullet("Platform dagilimi: Android=193, iOS=128")

H("1.3 Ham Veri ile On-Islenmis Verinin Karsilastirilmasi", 2)
P("On isleme (kucuk harfe cevirme, noktalama temizligi, stop-word cikarma, lemmatization / "
  "stemming) sonrasinda metinler belirgin sekilde sadelesmistir. Asagidaki tablo ham veri ile "
  "model egitiminde kullanilan iki temiz veri setinin boyut ve yapi karsilastirmasini ozetler:")
comp = pd.DataFrame({
    "Veri Seti": ["Ham (review)", "Lemmatized", "Stemmed"],
    "Satir": [321, 321, 321],
    "Toplam Kelime": [3260, 1951, 1951],
    "Ort. Kelime/Yorum": [10.16, 6.08, 6.08],
    "Benzersiz Kelime (Sozluk)": [891, 501, 433],
})
df_to_table(comp, font_size=9)
P("Yorum: On isleme ile toplam kelime sayisi 3260'tan 1951'e (~%40 azalma) inmistir; bu, stop-word "
  "ve noktalama temizliginin etkisidir. Benzersiz kelime sayisi ham veride 891 iken, lemmatized'de "
  "501, stemmed'de 433'e dusmustur. Stemming kelimeleri koklerine indirgerken farkli ekli "
  "varyantlari (ornegin 'sharing', 'shared', 'share' -> 'share') tek bicime topladigi icin "
  "sozluk lemmatization'a gore daha da kuculmustur. Daha kucuk ve yogun bir sozluk, kucuk veri "
  "setinde her kelime icin daha fazla baglamsal ornek demektir; bu da vektorlerin daha kararli "
  "ogrenilmesine yardimci olabilir.")
P("Not: Word2Vec egitiminde min_count=2 esigi uygulandigi icin yalnizca 2 ve daha fazla gecen "
  "kelimeler modele alinmistir. Bu nedenle modellerin nihai kelime hazinesi lemmatized icin 270, "
  "stemmed icin 243 kelimedir.")

H("1.4 GitHub Reposu", 2)
P("Tum calisma kodlari (model egitimi, benzerlik hesaplama, degerlendirme ve rapor uretimi) ve "
  "egitilen 16 model GitHub reposuna eklenmistir. Modeller .model uzantili olarak 'odev2/model' "
  "klasorunde, kodlar ise repo kokunde yer almaktadir. Calistirma talimatlari README dosyasinda "
  "verilmistir.")
bullet("Repo: https://github.com/miracbirbenhub/YAPAYZEKAPROJE")

doc.add_page_break()

# ======================================================================================
# 2. YONTEM
# ======================================================================================
H("2. Yontem", 1)

H("2.1 Word2Vec Vektorlestirme (Gorev-1)", 2)
P("Gensim kutuphanesinin Word2Vec sinifi kullanilarak, her iki veri seti icin asagidaki 8 "
  "parametre setiyle ayri ayri model egitilmistir (toplam 16 model). CBOW modelleri sg=0, "
  "SkipGram modelleri sg=1 ile egitilmistir. Kucuk veri setinde vektorlerin yeterince ogrenilmesi "
  "icin epochs=60, gurultu azaltmak icin min_count=2 ve tekrarlanabilirlik icin workers=1, "
  "seed=42 secilmistir.")
bullet("Ortak parametreler: min_count=2, epochs=60, workers=1, seed=42")
bullet("Degisen parametreler: model_type (cbow/skipgram), window (2/4), vector_size (100/300)")

P("Egitilen 16 modelin adlandirmasi ve yapilandirmasi:", bold=True)
model_tbl = cos[["model"]].copy()
model_tbl["veri_seti"] = model_tbl["model"].apply(lambda s: s.split("_")[1])
model_tbl["tur"] = model_tbl["model"].apply(lambda s: s.split("_")[2])
model_tbl["window"] = model_tbl["model"].apply(lambda s: s.split("_")[3].replace("win", ""))
model_tbl["boyut"] = model_tbl["model"].apply(lambda s: s.split("_")[4].replace("dim", ""))
model_tbl.columns = ["Model Adi", "Veri Seti", "Tur", "Window", "Vektor Boyutu"]
df_to_table(model_tbl, font_size=8)

P("Gorev-1 - Ornek Vektor Ciktilari (Anlamli kelime: 'premium')", bold=True)
P("Her model icin secilen onemli kelimenin ('premium') vektor uzayindaki en yakin 5 kelimesi "
  "asagida verilmistir. Amac, modelden modele benzerlik skorlarinin ve komsuluk iliskilerinin "
  "nasil degistigini gostermektir. (Bu skorlar tek basina modelin basarisini olcmek icin yeterli "
  "degildir; yalnizca anlamsal komsuluk hakkinda fikir verir.)")
df_to_table(sw.rename(columns={"model": "Model", "anahtar_kelime": "Kelime",
                               "en_yakin_5_kelime": "En Yakin 5 Kelime (skor)"}), font_size=7)
P("Gozlem: Tum modellerde 'premium' kelimesinin en yakin komsulari arasinda 'subscription', "
  "'worth', 'helpful', 'unlocks', 'extra' gibi anlamsal olarak gercekten ilgili kelimeler yer "
  "almaktadir; bu da modellerin temadaki kelime iliskilerini basariyla yakaladigini gosterir. "
  "SkipGram modellerinde skorlar (0.97-0.99) CBOW'a (0.99+) gore biraz daha dusuk ve daha "
  "ayristiricidir; cunku SkipGram nadir kelimeleri ve ince anlam farklarini daha iyi ogrenir. "
  "Beklentimiz: SkipGram modellerinin daha ayristirici (daha anlamli skor dagilimina sahip) "
  "vektorler uretmesi, CBOW'un ise daha hizli ve genel temada tutarli sonuc vermesidir.")

H("2.2 Benzerlik Hesaplama Yontemi (Gorev-2)", 2)
P("Secilen ornek giris metni veri setinin kendi icinden alinmistir (disaridan veri "
  "kullanilmamistir). Giris metni, premium fiyat temasini temsil eden net bir yorumdur:")
for line in query_txt:
    P("   " + line, italic=True)
P("Benzerlik su sekilde hesaplanmistir:")
bullet("Hem giris metni hem de veri setindeki her yorum icin, metindeki kelimelerin model "
       "vektorlerinin ARITMETIK ORTALAMASI alinarak bir dokuman vektoru olusturulur.")
bullet("Modelde karsiligi olmayan (OOV) kelimeler atlanir. Bir metindeki hicbir kelime modelde "
       "yoksa, NaN / ZeroDivisionError hatasini onlemek icin tamamen sifirlardan olusan SIFIR "
       "VEKTORU atanir (savunma mekanizmasi).")
bullet("Giris metni vektoru ile her dokuman vektoru arasinda COSINE SIMILARITY hesaplanir; en "
       "yuksek skorlu ilk 5 dokuman (giris metninin kendisi haric) o model icin sonuc olarak "
       "secilir.")
bullet("Bu islem 16 modelin her biri icin ayri ayri yapilmistir: 1 model icin 5, 16 model icin "
       "toplam 80 benzer dokuman cikarilmistir.")

doc.add_page_break()

# ======================================================================================
# 3. SONUCLAR VE DEGERLENDIRME
# ======================================================================================
H("3. Sonuclar ve Degerlendirme", 1)

H("3.1 Her Model icin Ilk 5 Benzer Metin", 2)
P("Asagida, secilen giris metnine her modelin getirdigi en benzer 5 yorum (ham metin haliyle), "
  "cosine skoru ve elle verilen anlamsal puan (1-5) ile birlikte listelenmistir.")
for model in cos["model"].tolist():
    sub = top5[top5["model"] == model].sort_values("rank")
    P(model, bold=True, size=10)
    tb = sub[["rank", "cosine", "anlamsal_1_5", "yorum_metni"]].copy()
    tb.columns = ["Sira", "Cosine", "Anlamsal(1-5)", "Yorum Metni"]
    df_to_table(tb, col_widths=[0.4, 0.7, 0.9, 4.5], font_size=8)
    doc.add_paragraph()

doc.add_page_break()

H("3.2 (1) Cosine Degerlendirme (Objective Evaluation)", 2)
P("Her model icin ilk 5 sonucun cosine benzerlik skorlari ve ortalamalari:")
cos_tbl = cos.copy()
cos_tbl.columns = ["Model", "Skor1", "Skor2", "Skor3", "Skor4", "Skor5", "Ortalama"]
df_to_table(cos_tbl, font_size=8)
P("Yorum: Tum modellerin ortalama cosine skorlari oldukca yuksektir (~0.997 - 0.9999). Bunun "
  "nedeni, dokuman vektorlerinin kisa yorumlardaki kelime vektorlerinin ortalamasi olarak "
  "hesaplanmasi ve kucuk/temadasal yogun bir korpusta tum yorum vektorlerinin birbirine cok yakin "
  "yonlerde olusmasidir. Bu nedenle MUTLAK cosine degeri tek basina basari icin zayif bir "
  "ayristiricidir - nitekim odev metninde de bu skorlarin tek basina yeterli olmadigi "
  "belirtilmistir. Yine de SIRALAMA dogru calismaktadir: getirilen yorumlarin neredeyse tamami "
  "premium/fiyat temasindadir. Ayristirma gucu acisindan SkipGram modelleri (ozellikle "
  "skipgram_win4_dim100, ort. ~0.997) CBOW'a gore biraz daha genis bir skor araligi uretmistir.")

H("3.3 (2) Anlamsal Degerlendirme (Subjective Evaluation)", 2)
P("Her modelin onerdigi 5 metne, giris metniyle anlamsal yakinligina gore 1-5 arasi elle puan "
  "verilmistir (1: cok alakasiz ... 5: neredeyse ayni temada, cok guclu benzerlik). Asagida model "
  "basina 5 puan ve aritmetik ortalamalari verilmistir:")
sem_tbl = sem.copy()
sem_tbl.columns = ["Model", "P1", "P2", "P3", "P4", "P5", "Ortalama"]
df_to_table(sem_tbl, font_size=8)
best_sem = sem.sort_values("ortalama_anlamsal", ascending=False).iloc[0]
P("Yorum: Anlamsal ortalamalar 4.6 - 4.8 araliginda yogunlasmaktadir; bu, modellerin getirdigi "
  "ilk 5 sonucun gercekten de giris metniyle ayni temada (premium fiyat / abonelik / deger) "
  "oldugunu dogrular. En yuksek anlamsal ortalamayi alan modeller cogunlukla CBOW + window=2 "
  "yapilandirmalaridir (ornegin word2vec_lemmatized_cbow_win2_dim100 ve "
  "word2vec_lemmatized_cbow_win2_dim300, ort. 4.8). Bu modeller, dar pencere sayesinde kelimenin "
  "yakin baglamina odaklanip tema tutarliligini iyi yakalamistir. Model yapilandirmalarinin "
  "(CBOW/SkipGram, window, vektor boyutu) etkisi gozlemlenebilmektedir: window=2 modelleri "
  "window=4'e gore biraz daha yuksek anlamsal puan almistir; vektor boyutunun (100 vs 300) anlamsal "
  "puana etkisi ise bu kucuk veri setinde sinirlidir.")

H("3.4 (3) Siralama Tutarliligi - Jaccard (Ranking Agreement)", 2)
P("Modellerin ilk 5 sonuc listelerinin ortusmesi Jaccard benzerligi ile olculmus ve 16x16'lik bir "
  "matris olusturulmustur. Jaccard, iki modelin AYNI giris metnine verdigi sonuclarin ne kadar "
  "ortustugunu olcer (tekil basari degil, modeller arasi tutarlilik). Kosegen, her modelin kendisiyle "
  "kiyasi oldugu icin 1.00'dir ve yorumlamaya katilmaz. Asagidaki heatmap matrisi gorsellestirir:")
img = os.path.join(OUT, "jaccard_heatmap.png")
doc.add_picture(img, width=Inches(6.3))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
P("Yorum: Off-diagonal Jaccard degerleri 0.25 ile 1.00 arasinda degismekte, ortalama ~0.45'tir. "
  "Heatmap'te belirgin bir oruntu vardir: ayni TUR (CBOW veya SkipGram) ve ayni window degerine "
  "sahip modeller birbirine cok daha benzer sonuc listeleri uretmektedir (yuksek Jaccard, acik "
  "renkler). Ozellikle CBOW modelleri kendi aralarinda; SkipGram modelleri de kendi aralarinda "
  "kume olusturmaktadir. Ayrica lemmatized ve stemmed setlerinin ayni yapilandirmali modelleri de "
  "yer yer yuksek ortusme gostermektedir; cunku stemming ve lemmatization ayni yorumlari benzer "
  "sekilde sadelestirir. Birbirine en cok benzeyen modeller ayni (tur+window) ailesinden cikmistir; "
  "bu, model yapilandirmasinin siralama davranisini dogrudan etkiledigini gosterir.")

H("3.5 Genel Karsilastirma ve En Basarili Modeller", 2)
P("Ozet tablo (model basina ortalama cosine ve ortalama anlamsal puan):", bold=True)
sum_tbl = summary.copy()
sum_tbl.columns = ["Model", "Ort. Cosine", "Ort. Anlamsal"]
df_to_table(sum_tbl, font_size=8)
P("Anlamsal vs Cosine: Cosine skorlari neredeyse tum modellerde tavana yakindir (ayristirma gucu "
  "dusuk), ancak anlamsal degerlendirme modeller arasinda gercek farki ortaya koyar. Iki olcut de "
  "ayni yonu gosterir: getirilen sonuclar tematik olarak dogrudur, fakat ince kalite farki "
  "(window=2 / CBOW lehine) yalnizca anlamsal puanda ve Jaccard kumelenmesinde net gorulur. Bu da "
  "neden tek bir olcute (sadece cosine) guvenmememiz gerektigini gosterir.")
P("Degerlendirme (en basarili / orta / en basarisiz):", bold=True)
bullet("EN BASARILI: CBOW + window=2 modelleri (word2vec_lemmatized_cbow_win2_dim100/dim300, "
       "word2vec_stemmed_cbow_win2_dim100). Hem en yuksek anlamsal ortalamayi (4.8) aldilar hem de "
       "tutarli, tema-odakli ilk 5 sonuc urettiler. Dar pencere, kisa yorumlarda kelimenin yakin "
       "baglamini iyi modelledigi icin basarilidir.")
bullet("ORTA: SkipGram + window=2 ve CBOW + window=4 modelleri. Sonuclari hala tematik olarak "
       "dogru; anlamsal ortalamalari 4.6 civarindadir.")
bullet("GORECELI EN ZAYIF: SkipGram + window=4 modelleri (skipgram_win4_dim100, ort. ~0.997 cosine, "
       "biraz daha dagilik sonuc). Genis pencere + SkipGram, kucuk korpusta daha gurultulu "
       "komsuluklar urettigi icin siralama tutarliligi biraz dusmustur.")
P("Neden bu modeller basarili bulundu? Veri setimiz kisa, tek dilli ve tematik olarak yogun "
  "(uygulama yorumlari) oldugu icin, kelimenin yakin baglamina odaklanan CBOW + dar pencere "
  "yaklasimi en kararli ve tema-tutarli vektorleri uretmistir. SkipGram daha buyuk ve seyrek "
  "verilerde one cikan bir yontemdir; bu olcekte avantaji tam yansimamistir.")

H("3.6 Model Yapilandirmalarinin Basariya Etkisi", 2)
bullet("Tur (CBOW vs SkipGram): CBOW bu kucuk korpusta biraz daha yuksek anlamsal puan ve daha "
       "tutarli siralama; SkipGram daha ayristirici ama biraz daha gurultulu skorlar verdi.")
bullet("Window (2 vs 4): window=2 (dar pencere) bu kisa metinlerde window=4'e gore biraz daha "
       "iyi tema tutarliligi sagladi.")
bullet("Vektor boyutu (100 vs 300): Bu veri olceginde 100 ile 300 arasinda anlamli bir kalite "
       "farki gozlenmedi; 300 boyut daha buyuk model dosyasi getirmesine ragmen kucuk korpusta ek "
       "fayda saglamadi (overfitting/seyreklik riski).")

doc.add_page_break()

# ======================================================================================
# 4. SONUC VE ONERILER
# ======================================================================================
H("4. Sonuc ve Oneriler", 1)
P("Bu calismada, Odev-1'de hazirladigimiz lemmatized ve stemmed veri setleri uzerinde 16 Word2Vec "
  "modeli egitilmis; secilen bir giris yorumuna en benzer 5 yorum her model icin bulunmus ve "
  "modeller cosine, anlamsal ve Jaccard olcutleriyle karsilastirilmistir. Temel cikarimlar:")
bullet("Kisa ve tematik yogun yorum verilerinde, CBOW + dar pencere (window=2) yapilandirmasi en "
       "tutarli ve anlamli benzerlik sonuclarini uretti; bu nedenle benzer kisa-metin/yorum "
       "benzerligi gorevleri icin CBOW+win2 onerilir.")
bullet("Mutlak cosine skoru tek basina yaniltici olabilir (tum modellerde tavana yakin). Model "
       "kalitesini olcerken cosine'i mutlaka anlamsal degerlendirme ve Jaccard tutarliligi ile "
       "birlikte yorumlamak gerekir.")
bullet("SkipGram ve buyuk vektor boyutu (300) bu olcekte ek fayda saglamadi; bunlar daha buyuk ve "
       "cesitli korpuslarda (ornegin tum platformdan on binlerce yorum) one cikabilir. Veri "
       "buyutuldukce SkipGram+win4+dim300 yeniden denenmelidir.")
bullet("Stemming, lemmatization'a gore daha kucuk bir sozluk uretti; her ikisi de benzer benzerlik "
       "sonuclari verdi. Anlamsal okunabilirlik onemliyse lemmatized, sozluk kucukluğu/hiz onemliyse "
       "stemmed tercih edilebilir.")
P("Hangi model hangi gorev icin? Hizli ve tema-tutarli yorum/oneri benzerligi icin "
  "CBOW+window2+dim100 (kucuk, hizli, basarili); ince anlam ayrimi gerektiren ve veri bollugu olan "
  "senaryolar icin SkipGram+window4+dim300 onerilir.")

out_docx = os.path.join(REP, "Odev2_Raporu.docx")
doc.save(out_docx)
print("DOCX yazildi:", out_docx)
